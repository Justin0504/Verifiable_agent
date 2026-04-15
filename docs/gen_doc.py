# -*- coding: utf-8 -*-
"""Generate a well-formatted .docx for the Verifiable Agent architecture."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Global defaults ──
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
rPr.append(rFonts)

for level, (sz, color) in {
    1: (22, "1A1A2E"),
    2: (16, "16213E"),
    3: (13, "0F3460"),
}.items():
    h = doc.styles["Heading %d" % level]
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor(*bytes.fromhex(color))
    h.font.bold = True
    h.font.name = "Microsoft YaHei"
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)
    rPr2 = h.element.get_or_add_rPr()
    rF2 = OxmlElement("w:rFonts")
    rF2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rF2)


def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p


def add_rich(segments, space_after=None):
    p = doc.add_paragraph()
    for text, fmt in segments:
        run = p.add_run(text)
        run.bold = fmt.get("b", False)
        run.italic = fmt.get("i", False)
        if "s" in fmt: run.font.size = Pt(fmt["s"])
        if "c" in fmt: run.font.color.rgb = RGBColor(*bytes.fromhex(fmt["c"]))
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_head=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2)
    if bold_head:
        p.add_run(bold_head).bold = True
    p.add_run(text)
    return p


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def make_table(headers, rows, col_widths=None, header_bg="1A1A2E"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        bg = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tbl.rows[1 + ri].cells[ci]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            set_cell_bg(c, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


# === Chinese content stored safely ===
# Using plain ASCII quotes inside strings; curly quotes as \u escapes

Q = "\u201c"  # left curly "
QQ = "\u201d"  # right curly "
AR = "\u2192"  # →
DA = "\u2014"  # —

# ═══════════════════════════════════════════════
#  TITLE
# ═══════════════════════════════════════════════
add_para("Verifiable Agent", bold=True, size=28, color="1A1A2E",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("自演化幻觉检测框架 " + DA + " 架构与实现细节",
         size=14, color="666666", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

# ═══════════════════════════════════════════════
doc.add_heading("一、我们在解决什么问题", level=1)

add_para(
    "大语言模型（LLM）会生成看起来很合理但实际上不正确的内容，也就是" + Q + "幻觉" + QQ + "。"
    "现有的检测方法（FActScore、SAFE、CoVe 等）都是静态的" + DA + "用固定的 prompt 跑一遍就完了，"
    "不会从自己的错误中学习。"
)
add_para(
    "我们的方案不一样：搭一个能自我进化的检测系统。不微调模型参数，"
    "纯靠优化 prompt 和积累推理规则，让检测准确率随训练轮次持续上升。"
    "核心思路是" + Q + "出题" + AR + "验证" + AR + "找错" + AR + "改进" + AR + "再出题" + QQ + "的闭环迭代。",
    space_after=12
)

# ═══════════════════════════════════════════════
doc.add_heading("二、系统整体架构", level=1)

add_para(
    "系统由三个阶段组成一个流水线（Proposer " + AR + " Responder " + AR + " Verifier），"
    "加上一个自演化循环把它们串成闭环。"
    "每个 epoch（训练轮次）结束后，系统的五个核心组件全部更新：",
    space_after=8
)

make_table(
    ["组件", "名称", "作用", "怎么更新"],
    [
        ["P_t", "Proposer 策略", "控制生成什么样的训练数据", "根据弱点画像调整出题方向"],
        ["V_t", "Verifier Prompt", "验证时用的 system prompt", "通过 prompt 精炼 + 验证门控更新"],
        ["R_t", "ReasoningBank", "结构化推理规则库", "从成功/失败轨迹中蒸馏新规则"],
        ["M_t", "弱点画像", "Verifier 当前的错误模式", "每轮根据校准集结果更新"],
        ["D_t", "训练数据集", "经过筛选的高质量验证数据", "每轮重新生成 + 过滤"],
    ],
    col_widths=[2, 3.5, 5, 5.5],
)

# ═══════════════════════════════════════════════
doc.add_heading("三、三阶段流水线详解", level=1)

# --- Stage 1 ---
doc.add_heading("Stage 1：Proposer（出题者）", level=2)

add_para(
    "Proposer 负责生成高质量的 claim-evidence 训练对。"
    "这里有两个关键设计："
)

add_rich([
    ("设计一：验证技能导向。", {"b": True}),
    ("我们不是让 LLM 随便编数据，而是要求它按照六类具体的验证场景来出题。"
     "每一类测试的是不同的推理技能，不是知识储备。", {}),
], space_after=6)

make_table(
    ["验证场景", "标签", "测什么能力", "具体例子"],
    [
        ["改写识别", "S", "能否看出不同措辞说的是同一件事",
         "证据：" + Q + "拿破仑身高5尺7寸" + QQ + "\n声明：" + Q + "拿破仑身高在当时算中等水平" + QQ],
        ["微妙矛盾", "C", "能否抓住一个关键细节的变化",
         "证据：" + Q + "该研究于2019年发表" + QQ + "\n声明：" + Q + "该研究于2018年发表" + QQ],
        ["话题重叠", "N", "能否区分同领域但不同事实",
         "证据：讲的是火星大气成分\n声明：讲的是火星地表温度"],
        ["部分匹配", "S", "能否接受精度较低但正确的表述",
         "证据：" + Q + "GDP增长了23.5%" + QQ + "\n声明：" + Q + "GDP增长超过20%" + QQ],
        ["否定翻转", "C", "能否识别语义取反",
         "证据：" + Q + "没有证据表明X有效" + QQ + "\n声明：" + Q + "研究证实X有效" + QQ],
        ["无关内容", "N", "能否判断完全不相关的话题",
         "证据：讲量子物理\n声明：讲法国大革命"],
    ],
    col_widths=[2.5, 1.2, 4, 8],
)

add_rich([
    ("设计二：弱点感知生成。", {"b": True}),
    ("Proposer 不是均匀地出题，而是根据 Verifier 上一轮暴露的弱点来调整。"
     "比如上一轮发现 S" + AR + "C 误判最多（12次），下一轮就会重点出" + Q + "改写识别" + QQ + "和" + Q + "部分匹配" + QQ
     + "这类容易被误判为矛盾的题目。同时还会专门生成对抗性样本" + DA, {}),
    ("刻意让 claim 和 evidence 看起来很像但标签不同", {"b": True, "c": "C62828"}),
    ("，逼 Verifier 学会精细区分。", {}),
], space_after=10)

# --- Stage 2 ---
doc.add_heading("Stage 2：Responder（被测模型）", level=2)

add_para("Responder 是被检测的目标 LLM。在两种场景下有不同的角色：")
add_bullet("直接使用数据集提供的回答。比如 TruthfulQA 里每个问题已经有正确答案和常见错误答案，我们直接拿来验证。",
           bold_head="Benchmark 评估：")
add_bullet("接入任意目标 LLM（GPT-4o、Claude、Llama-3 等），向它发送 Proposer 生成的问题，收集真实回复。",
           bold_head="实际部署：")
add_para("", space_after=6)

# --- Stage 3 ---
doc.add_heading("Stage 3：Verifier（验证者）", level=2)

add_para("Verifier 是整个系统的核心，完成从原始回复到幻觉评分的全过程。分四个子步骤：")

doc.add_heading("3.1 原子声明分解", level=3)
add_para(
    "用 LLM 把模型回复拆解成独立的、可单独验证的事实声明。"
    "每个声明只包含一个可验证的事实断言。"
)
add_rich([
    ("示例输入：", {"b": True}),
    (Q + "拿破仑是法国皇帝，他在1815年的滑铁卢战役中败给了英国联军。" + QQ, {"i": True}),
])
add_rich([
    ("示例输出：", {"b": True}),
    ("[" + Q + "拿破仑是法国皇帝" + QQ + ", " + Q + "拿破仑在1815年的滑铁卢战役中战败" + QQ
     + ", " + Q + "拿破仑的对手是英国联军" + QQ + "]", {"i": True}),
], space_after=8)

doc.add_heading("3.2 证据检索", level=3)
add_para(
    "为每个原子声明从知识库中检索最相关的证据段落。"
    "在 benchmark 评估中，证据直接由数据集提供（比如 TruthfulQA 的 best_answer + correct_answers）。"
    "在通用场景中可以对接 Wikipedia、搜索引擎等外部知识源。",
    space_after=8
)

doc.add_heading("3.3 声明级验证（S/C/N 分类）", level=3)
add_para("对每个声明-证据对，LLM 输出一个三分类判断和置信度：")

make_table(
    ["标签", "含义", "判断标准", "幻觉分"],
    [
        ["S (Supported)", "证据支持", "证据中能找到直接或间接支持该声明的信息", "0"],
        ["C (Contradicted)", "证据矛盾", "证据中的信息与该声明存在事实冲突", "1"],
        ["N (Not enough info)", "信息不足", "证据没有涉及该声明的内容，无法判断", "0.5"],
    ],
    col_widths=[3, 2, 7.5, 1.8],
)

add_para(
    "这一步是自演化重点优化的对象。Verifier 的 system prompt（V_t）决定了分类质量，"
    "而这个 prompt 会在每轮 epoch 中被精炼。"
    "如果 ReasoningBank 中有和当前声明相关的推理规则，也会注入到 prompt 中辅助判断。",
    space_after=8
)

doc.add_heading("3.4 置信度加权聚合", level=3)
add_para("一个回复可能拆出 3-6 个声明，需要汇总成一个回复级别的判断。我们的聚合规则：")
add_bullet("如果有任一声明被高置信度（>= 0.8）判为 C，整个回复判为 C（矛盾优先）")
add_bullet("否则，按每个声明的置信度做加权投票，票数最高的标签胜出")
add_bullet("平票时优先级：C > N > S（保守策略，倾向于报告潜在问题）")

add_rich([
    ("为什么不用旧的严格规则", {"b": True}),
    ("（" + Q + "只要有一个 C 就判 C，全是 S 才判 S" + QQ + "）？"
     "因为 TruthfulQA 的 S 类样本有 3-6 个 claims，只要其中一个被误判为 N，"
     "整个样本就会从 S 变成 N。这是之前 Verifier 只有 59% 而 FActScore 有 71% 的主要原因。", {}),
], space_after=12)

# ═══════════════════════════════════════════════
doc.add_heading("四、自演化循环（核心创新）", level=1)

add_para(
    "自演化是区别于所有 baseline 方法的核心。"
    "每轮 epoch 执行以下完整流程，形成" + Q + "出错" + AR + "学习" + AR + "变强" + QQ + "的闭环：",
    space_after=10
)

# Step 0
doc.add_heading("Step 0：基线校准（仅首轮）", level=2)
add_para(
    "从 benchmark 数据中取出 20% 作为校准集（calibration set），剩余 80% 作为最终测试集。"
    "用 epoch 0 的 Verifier（没有任何演化）跑一遍校准集，记录："
)
add_bullet("基线准确率（我们的实验中是 71.2%）")
add_bullet("错误分布：哪些标签之间最容易搞混（比如 S" + AR + "C 有 12 次，C" + AR + "N 有 8 次）")
add_bullet("初始弱点画像：告诉 Proposer 第一轮该重点出什么题")

add_rich([
    ("为什么用真实数据做校准？", {"b": True, "c": "C62828"}),
    ("之前的版本（v8）只在合成数据上评估，结果合成数据太简单，"
     "Verifier 全答对，学不到东西。用真实 benchmark 数据才能发现真正的弱点。", {}),
], space_after=10)

# Step 1
doc.add_heading("Step 1：生成训练数据", level=2)
add_para("每轮 epoch 生成两类数据：")

make_table(
    ["数据类型", "每轮数量", "生成方式", "目的"],
    [
        ["验证技能数据", "约 24 条", "LLM 按六类场景生成", "覆盖各种验证能力"],
        ["对抗性数据", "约 12 条", "LLM 根据弱点画像定向生成", "专攻 Verifier 的薄弱环节"],
    ],
    col_widths=[3, 2.5, 5, 5],
)

# Step 2
doc.add_heading("Step 2：金标签验证（Self-Challenging）", level=2)
add_para(
    "LLM 生成的数据质量参差不齐，标签可能是错的。"
    "我们让另一个 LLM 调用独立检查每个样本的标签是否正确："
)
add_bullet("给 LLM 看 claim、evidence、标签，问" + Q + "这个标签对不对？为什么？" + QQ)
add_bullet("如果判定标签不对，直接丢弃这个样本")
add_bullet("通常保留 80-90% 的样本（比 Self-Challenging 论文中 5% 的留存率宽松，因为我们的生成质量更高）")

add_rich([
    ("核心思想：", {"b": True}),
    ("宁可训练数据少，也要保证每一条的标签是对的。"
     "用错误标签训练只会让 Verifier 越来越差。", {}),
], space_after=10)

# Step 3
doc.add_heading("Step 3：Soft R-Zero 数据筛选", level=2)
add_para(
    "不是所有数据都有同等的训练价值。"
    "Verifier 已经能轻松答对的题没有学习意义，真正有价值的是" + Q + "边界案例" + QQ
    + DA + DA + "Verifier 半对半错、很纠结的那些。"
)
add_para("具体做法：对每个样本让 Verifier 独立判断 3 次，统计正确率：")

make_table(
    ["3 次判断结果", "正确率", "不确定性", "处理"],
    [
        ["全对 (3/3)", "100%", "低 " + DA + " 太简单了", "排在后面，可能被丢弃"],
        ["对 2 错 1 (2/3)", "67%", "中等 " + DA + " 有点难", "保留"],
        ["对 1 错 2 (1/3)", "33%", "高 " + DA + " 很纠结", "优先保留"],
        ["全错 (0/3)", "0%", "低 " + DA + " 可能标签有问题", "排在后面，可能被丢弃"],
    ],
    col_widths=[3, 2, 4, 5.5],
)

add_para("按不确定性从高到低排序，保留前 60%。")
add_rich([
    ("为什么叫 Soft R-Zero？", {"b": True}),
    ("原始 R-Zero 论文用硬阈值过滤（正确率在 25%-75% 之间才保留），"
     "但我们发现简单的合成数据上 Verifier 要么全对要么全错，硬过滤会把数据全删光"
     "（v8 就是这样挂的，52 条全删到 0）。"
     "所以我们改成排序 + 取 top-60%，保证每轮至少有足够的训练数据。", {}),
], space_after=10)

# Step 4
doc.add_heading("Step 4：训练 + Prompt 精炼 + 验证门控", level=2)
add_para("这是自演化最核心的一步，分三个阶段：")

add_rich([
    ("阶段 A：在合成数据上跑 Verifier。", {"b": True}),
    ("记录每个样本的预测结果，分析错误分布（哪些标签对混淆最严重、具体错在哪里）。", {}),
], space_after=4)

add_rich([
    ("阶段 B：LLM 生成改进后的 prompt。", {"b": True}),
    ("把当前 prompt、错误分析、具体错误案例发给 LLM，让它写一个更好的 system prompt。"
     "要求新 prompt 控制在 250 词以内，明确 S/C/N 的判断标准，针对性修复错误模式。", {}),
], space_after=4)

add_rich([
    ("阶段 C：在真实校准集上验证。", {"b": True, "c": "C62828"}),
    ("这是关键" + DA + DA + "新 prompt 不是生成出来就用，必须在真实 benchmark 校准集上的准确率", {}),
    ("严格高于", {"b": True, "c": "C62828"}),
    ("当前最优 prompt 才会被接受。如果没有提升甚至变差了，直接回滚，当这轮 prompt 改动没发生过。", {}),
])

add_para("")
make_table(
    ["情况", "操作", "原因"],
    [
        ["新 prompt 校准准确率 > 旧", "接受，更新最优 prompt", "确实变强了"],
        ["新 prompt 校准准确率 = 旧", "拒绝，回滚", "没有提升不算进步"],
        ["新 prompt 校准准确率 < 旧", "拒绝，回滚", "变差了，不能用"],
    ],
    col_widths=[5.5, 4.5, 5],
)

# Step 5
doc.add_heading("Step 5：ReasoningBank 规则蒸馏", level=2)
add_para("从每轮的预测结果中提取可复用的推理规则，存入 ReasoningBank：")

make_table(
    ["规则来源", "每轮提取", "内容", "作用"],
    [
        ["失败案例", "最多 3 条", Q + "别这么做" + QQ + "的反面教训",
         "比如" + Q + "当 claim 和 evidence 讲的是同一实体的不同属性时，应该判 N 而不是 C" + QQ],
        ["成功案例", "最多 2 条", Q + "这么做有效" + QQ + "的正面经验",
         "比如" + Q + "当 claim 用不同数字表述但量级一致时，应该判 S" + QQ],
    ],
    col_widths=[2.5, 2.2, 4, 7],
)

add_para(
    "规则的使用方式：下次 Verifier 遇到新的声明时，从 ReasoningBank 中检索最相关的 1 条规则，"
    "注入到当前的 prompt 中辅助判断。只取 top-1 是为了避免 prompt 膨胀" + DA + DA
    + "塞太多规则反而会干扰 LLM 的判断。",
    space_after=10
)

# Step 6
doc.add_heading("Step 6：弱点画像更新，驱动下一轮", level=2)
add_para("根据校准集上的最新错误分布，更新 Proposer 的出题指导：")
add_bullet("继续多出改写识别和部分匹配的题",
           bold_head="S" + AR + "C 混淆仍然最严重（8 次）" + AR + " ")
add_bullet("这类题可以减少，已经学会了",
           bold_head="C" + AR + "N 从 8 次降到了 3 次 " + AR + " ")
add_bullet("下一轮重点出话题重叠的题",
           bold_head="N" + AR + "S 是新出现的问题（4 次）" + AR + " ")

add_para(
    "这样形成了一个完整的闭环：弱点画像指导 Proposer 出题 " + AR
    + " Verifier 在难题上训练 " + AR
    + " prompt 精炼修复弱点 " + AR
    + " 新的弱点画像 " + AR
    + " 下一轮......",
    space_after=12
)

# ═══════════════════════════════════════════════
doc.add_heading("五、测试阶段", level=1)
add_para("15 轮 epoch 训练结束后，用最优 prompt 在独立测试集（80%）上做最终评估：")
add_bullet("每个样本跑 3 次独立验证，取多数投票作为最终预测（R-Zero 的测试时集成思路）")
add_bullet("与 6 个 baseline 方法在完全相同的测试集上对比，保证公平性")
add_bullet("报告 Accuracy、Macro F1、每类标签的混淆矩阵")

add_para("")

make_table(
    ["Baseline 方法", "类型", "核心思路"],
    [
        ["Direct Prompting", "零样本", "直接问 LLM 这个回答对不对"],
        ["FActScore", "分解+验证", "拆成原子事实逐个验证（我们 Stage 3 的基础版）"],
        ["SAFE", "多步搜索", "Search-Augmented Factual Evaluation，用搜索引擎辅助"],
        ["CoVe", "自我一致性", "Chain-of-Verification，让模型自己检查自己"],
        ["SelfCheckGPT", "采样对比", "多次采样同一问题，比较回答一致性"],
        ["Retrieve+NLI", "检索+推理", "检索相关文档 + NLI 模型判断蕴含/矛盾"],
    ],
    col_widths=[3, 2.5, 10],
)

# ═══════════════════════════════════════════════
doc.add_heading("六、参考论文及借鉴", level=1)

make_table(
    ["论文", "出处", "我们借鉴了什么", "怎么用的"],
    [
        ["R-Zero", "ICLR 2026\n(2508.05004)",
         "不确定性驱动的数据选择",
         "Soft R-Zero 筛选 + 测试时 3 次投票"],
        ["Self-Challenging", "arXiv\n2506.01716",
         "对生成数据做可执行验证",
         "金标签验证，丢弃标签错误的样本"],
        ["ReasoningBank", "arXiv\n2509.25140",
         "从轨迹中蒸馏结构化规则",
         "成功+失败双向蒸馏，top-1 检索注入"],
        ["A-MEM", "arXiv\n2502.12110",
         "结构化记忆的设计思路",
         "弱点画像的存储和演化"],
    ],
    col_widths=[2.8, 2.5, 4.5, 5.5],
)

# ═══════════════════════════════════════════════
doc.add_heading("七、目前实验进展", level=1)

make_table(
    ["指标", "数值"],
    [
        ["评测数据集", "TruthfulQA (validation split, 400 样本)"],
        ["校准集大小", "80 样本（20%）"],
        ["测试集大小", "320 样本（80%）"],
        ["Epoch 0 校准准确率", "71.2%（无演化的 baseline）"],
        ["Epoch 1 校准准确率", "82.5%（+11.3 个百分点）"],
        ["当前状态", "15 轮 epoch 训练进行中"],
        ["对比 baseline 数量", "6 个"],
    ],
    col_widths=[4, 12],
)

add_para(
    "仅一轮自演化就把准确率从 71.2% 提升到 82.5%，验证了自演化机制的有效性。"
    "完整的 15 轮 epoch 训练和 baseline 对比结果出来后会更新本文档。",
    space_after=6
)

# ── Save ──
out = "/Users/justin/Verifiable_agent/docs/architecture_overview.docx"
doc.save(out)
print("Saved: " + out)
